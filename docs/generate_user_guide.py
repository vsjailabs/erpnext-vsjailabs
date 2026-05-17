#!/usr/bin/env python3
"""
Generate ERPNext End User Guide — Enterprise Grade DOCX
For: VSJ AI Labs
Author: Ashish Kumar Satyam
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import datetime

# ── Brand Colors ─────────────────────────────────────────────────────────────

NAVY       = RGBColor(0x1B, 0x3A, 0x5C)
BLUE       = RGBColor(0x2E, 0x75, 0xB6)
GREEN      = RGBColor(0x27, 0xAE, 0x60)
ORANGE     = RGBColor(0xE6, 0x7E, 0x22)
RED        = RGBColor(0xE7, 0x4C, 0x3C)
DARK       = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM     = RGBColor(0x7F, 0x8C, 0x8D)
WHITE      = RGBColor(0xFF, 0xFF, 0xFF)
LIGHT_BG   = RGBColor(0xF0, 0xF4, 0xF8)
ALT_ROW    = "F8F9FA"
WHITE_HEX  = "FFFFFF"
NAVY_HEX   = "1B3A5C"
BLUE_HEX   = "2E75B6"
GREEN_HEX  = "27AE60"
GREEN_BG   = "E8F5E9"
ORANGE_BG  = "FFF8E1"
RED_BG     = "FFEBEE"

TODAY = datetime.date.today().strftime("%B %d, %Y")
VERSION = "1.0"
DOC_TITLE = "ERPNext Portal — End User Guide"
DOC_SUBTITLE = "Role-Based Access & Operations Manual"
COMPANY = "VSJ AI Labs"
AUTHOR = "Ashish Kumar Satyam"
CLASSIFICATION = "Internal / Confidential"

# ── Helpers ──────────────────────────────────────────────────────────────────

def set_cell_shading(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_text(cell, text, bold=False, color=None, size=10, font="Arial", alignment=None):
    cell.text = ""
    p = cell.paragraphs[0]
    if alignment:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def add_header_row(table, texts, col_widths=None):
    row = table.rows[0]
    for i, text in enumerate(texts):
        set_cell_shading(row.cells[i], NAVY_HEX)
        set_cell_text(row.cells[i], text, bold=True, color=WHITE, size=10)

def add_data_row(table, texts, row_idx):
    row = table.add_row()
    bg = ALT_ROW if row_idx % 2 == 0 else WHITE_HEX
    for i, text in enumerate(texts):
        set_cell_shading(row.cells[i], bg)
        set_cell_text(row.cells[i], str(text), size=9)
    return row

def add_badge_row(table, texts, row_idx, badge_col=None, badge_color=None, badge_bg=None):
    row = table.add_row()
    bg = ALT_ROW if row_idx % 2 == 0 else WHITE_HEX
    for i, text in enumerate(texts):
        set_cell_shading(row.cells[i], bg)
        if i == badge_col and badge_color:
            set_cell_shading(row.cells[i], badge_bg or bg)
            cell = row.cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(f"● {text}")
            run.font.name = "Arial"
            run.font.size = Pt(9)
            run.font.bold = True
            run.font.color.rgb = badge_color
        else:
            set_cell_text(row.cells[i], str(text), size=9)
    return row

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Arial"
    return h

def add_body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.bold = bold
    return p

def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    p.clear()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    if level > 0:
        p.paragraph_format.left_indent = Cm(1.5 * level)
    return p

def add_note_box(doc, title, text):
    p = doc.add_paragraph()
    run = p.add_run(f"{title}: ")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = BLUE
    run2 = p.add_run(text)
    run2.font.name = "Arial"
    run2.font.size = Pt(10)
    run2.font.color.rgb = DARK

def make_table(doc, headers, rows, badge_col=None, badge_map=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    add_header_row(table, headers)
    for idx, row_data in enumerate(rows):
        if badge_col is not None and badge_map:
            val = row_data[badge_col]
            bc, bbg = badge_map.get(val, (DARK, WHITE_HEX))
            add_badge_row(table, row_data, idx, badge_col=badge_col, badge_color=bc, badge_bg=bbg)
        else:
            add_data_row(table, row_data, idx)
    return table

# ── Badge maps ───────────────────────────────────────────────────────────────

STATUS_BADGES = {
    "Full Access":    (GREEN,  GREEN_BG),
    "Read/Write":     (GREEN,  GREEN_BG),
    "Read Only":      (ORANGE, ORANGE_BG),
    "Limited":        (ORANGE, ORANGE_BG),
    "No Access":      (RED,    RED_BG),
    "Restricted":     (RED,    RED_BG),
    "Create/Edit":    (GREEN,  GREEN_BG),
    "Approve":        (BLUE,   GREEN_BG),
    "Submit":         (GREEN,  GREEN_BG),
    "View Only":      (ORANGE, ORANGE_BG),
}

# ── Main Document ────────────────────────────────────────────────────────────

def generate():
    doc = Document()

    # Page setup
    section = doc.sections[0]
    section.page_width = Cm(21.59)
    section.page_height = Cm(27.94)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    # Default font
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Arial"
    font.size = Pt(10)
    font.color.rgb = DARK

    # ═══════════════════════════════════════════════════════════════════════
    # COVER PAGE
    # ═══════════════════════════════════════════════════════════════════════

    for _ in range(6):
        doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(COMPANY)
    run.font.name = "Arial"
    run.font.size = Pt(32)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DOC_TITLE)
    run.font.name = "Arial"
    run.font.size = Pt(22)
    run.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(DOC_SUBTITLE)
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = MEDIUM

    doc.add_paragraph()
    doc.add_paragraph()

    # Cover metadata table
    meta_table = doc.add_table(rows=6, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Document Version", VERSION),
        ("Date", TODAY),
        ("Author", AUTHOR),
        ("For", COMPANY),
        ("Classification", CLASSIFICATION),
        ("ERPNext Version", "v15 (Latest Stable)"),
    ]
    for i, (key, val) in enumerate(meta_data):
        bg = ALT_ROW if i % 2 == 0 else WHITE_HEX
        set_cell_shading(meta_table.rows[i].cells[0], bg)
        set_cell_shading(meta_table.rows[i].cells[1], bg)
        set_cell_text(meta_table.rows[i].cells[0], key, bold=True, size=10)
        if key == "Classification":
            set_cell_text(meta_table.rows[i].cells[1], val, bold=True, color=RED, size=10)
        else:
            set_cell_text(meta_table.rows[i].cells[1], val, size=10)

    doc.add_paragraph()

    # Classification banner
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CLASSIFICATION)
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = RED

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # TABLE OF CONTENTS
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "Table of Contents", level=1)
    toc_items = [
        "1. Document Information",
        "2. Version History",
        "3. Introduction & Purpose",
        "4. System Access & Login",
        "5. Role Overview & Permissions Matrix",
        "6. System Administrator",
        "7. HR Manager",
        "8. HR User",
        "9. Accounts Manager",
        "10. Accounts User",
        "11. Sales Manager",
        "12. Sales User",
        "13. Purchase Manager",
        "14. Purchase User",
        "15. Stock Manager",
        "16. Stock User",
        "17. Manufacturing Manager",
        "18. Manufacturing User",
        "19. Projects Manager",
        "20. Projects User",
        "21. Support Team",
        "22. Website Manager",
        "23. Employee (Self-Service)",
        "24. Auditor (Read-Only)",
        "25. Dashboard & Reports Guide",
        "26. Common Workflows",
        "27. Security & Best Practices",
        "28. Troubleshooting & FAQ",
        "29. Support & Escalation",
        "30. Glossary",
    ]
    for item in toc_items:
        p = doc.add_paragraph()
        run = p.add_run(item)
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.color.rgb = DARK

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 1. DOCUMENT INFORMATION
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "1. Document Information", level=1)

    make_table(doc,
        ["Field", "Value"],
        [
            ["Document Name", DOC_TITLE],
            ["Version", VERSION],
            ["Date", TODAY],
            ["Owner", AUTHOR],
            ["Stakeholders", "All ERPNext Users, IT Admin, Department Heads"],
            ["Classification", CLASSIFICATION],
            ["ERPNext Version", "v15 (Latest Stable)"],
            ["Hosting", "Google Cloud Platform (Compute Engine)"],
            ["Portal URL", "http://34.121.203.2"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 2. VERSION HISTORY
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "2. Version History", level=1)

    make_table(doc,
        ["Version", "Date", "Author", "Changes"],
        [
            ["1.0", TODAY, AUTHOR, "Initial release — all roles documented"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 3. INTRODUCTION & PURPOSE
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "3. Introduction & Purpose", level=1)

    add_body(doc, "This document serves as the comprehensive end-user guide for the ERPNext portal deployed by VSJ AI Labs. It provides role-based instructions for all users to effectively navigate, operate, and manage their respective modules within ERPNext v15.")

    add_body(doc, "ERPNext is an open-source, full-featured Enterprise Resource Planning (ERP) system that integrates key business functions including Human Resources, Accounting, Sales, Purchasing, Inventory, Manufacturing, Projects, and Customer Support into a single unified platform.")

    add_heading(doc, "3.1 Who Should Read This Guide", level=2)
    add_bullet(doc, "System Administrators responsible for ERPNext configuration and user management")
    add_bullet(doc, "HR Managers and HR Users managing employee lifecycle and payroll")
    add_bullet(doc, "Accounts Managers and Users handling financial transactions and reporting")
    add_bullet(doc, "Sales and Purchase teams managing orders, quotations, and invoices")
    add_bullet(doc, "Stock/Inventory managers tracking warehouse operations")
    add_bullet(doc, "Manufacturing teams managing BOMs, work orders, and production planning")
    add_bullet(doc, "Project managers tracking tasks, timesheets, and project profitability")
    add_bullet(doc, "Support teams handling customer issues and service level agreements")
    add_bullet(doc, "Employees using self-service features (leave, expense, attendance)")
    add_bullet(doc, "Auditors with read-only access for compliance reviews")

    add_heading(doc, "3.2 ERPNext Modules Overview", level=2)

    make_table(doc,
        ["Module", "Description", "Key Documents"],
        [
            ["Human Resources", "Employee management, attendance, payroll, leave, recruitment", "Employee, Leave Application, Payroll Entry, Attendance"],
            ["Accounts", "General ledger, AP/AR, journal entries, tax, budgets", "Sales Invoice, Purchase Invoice, Payment Entry, Journal Entry"],
            ["Sales", "Lead-to-cash cycle, quotations, orders, delivery", "Lead, Quotation, Sales Order, Delivery Note, Sales Invoice"],
            ["Buying", "Procure-to-pay cycle, supplier management, RFQs", "Supplier, RFQ, Purchase Order, Purchase Receipt, Purchase Invoice"],
            ["Stock", "Inventory management, warehousing, serial/batch tracking", "Stock Entry, Stock Reconciliation, Delivery Note, Purchase Receipt"],
            ["Manufacturing", "Production planning, BOMs, work orders, quality", "BOM, Work Order, Job Card, Production Plan, Quality Inspection"],
            ["Projects", "Project tracking, tasks, timesheets, Gantt charts", "Project, Task, Timesheet, Activity Type"],
            ["Support", "Issue tracking, SLAs, knowledge base, warranties", "Issue, Warranty Claim, Service Level Agreement"],
            ["Website", "CMS, blog, web forms, portal pages", "Web Page, Blog Post, Web Form"],
            ["Setup", "Company, currency, email, print format, workflows", "Company, Workflow, Print Format, Naming Series"],
        ]
    )

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 4. SYSTEM ACCESS & LOGIN
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "4. System Access & Login", level=1)

    add_heading(doc, "4.1 Access URL", level=2)

    make_table(doc,
        ["Environment", "URL", "Status"],
        [
            ["Staging", "http://34.121.203.2", "Active"],
            ["Production", "TBD (custom domain pending)", "Planned"],
        ],
        badge_col=2,
        badge_map={"Active": (GREEN, GREEN_BG), "Planned": (ORANGE, ORANGE_BG)}
    )

    add_heading(doc, "4.2 Login Steps", level=2)
    add_body(doc, "Follow these steps to access the ERPNext portal:")
    add_bullet(doc, "Step 1: Open your web browser (Chrome, Firefox, Edge, or Safari recommended)")
    add_bullet(doc, "Step 2: Navigate to the portal URL: http://34.121.203.2")
    add_bullet(doc, "Step 3: Enter your Username (email address provided by your administrator)")
    add_bullet(doc, "Step 4: Enter your Password")
    add_bullet(doc, "Step 5: Click the 'Login' button")
    add_bullet(doc, "Step 6: On first login, you will be prompted to change your password")

    add_heading(doc, "4.3 Password Policy", level=2)
    add_bullet(doc, "Minimum 8 characters with at least 1 uppercase, 1 lowercase, 1 number, and 1 special character")
    add_bullet(doc, "Passwords expire every 90 days (configurable by administrator)")
    add_bullet(doc, "Account locks after 3 failed login attempts for 30 minutes")
    add_bullet(doc, "Use 'Forgot Password' link on the login page for self-service password reset")

    add_heading(doc, "4.4 Browser Requirements", level=2)
    make_table(doc,
        ["Browser", "Minimum Version", "Recommended"],
        [
            ["Google Chrome", "90+", "Latest"],
            ["Mozilla Firefox", "88+", "Latest"],
            ["Microsoft Edge", "90+", "Latest"],
            ["Apple Safari", "14+", "Latest"],
        ]
    )
    add_note_box(doc, "Note", "Internet Explorer is NOT supported. JavaScript must be enabled.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 5. ROLE OVERVIEW & PERMISSIONS MATRIX
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "5. Role Overview & Permissions Matrix", level=1)

    add_body(doc, "ERPNext uses a role-based access control (RBAC) model. Each user is assigned one or more roles that determine which modules, documents, and actions they can access. The following matrix provides a high-level overview of all roles and their access levels across modules.")

    add_heading(doc, "5.1 Role Summary", level=2)

    make_table(doc,
        ["Role", "Module", "Scope", "Typical Users"],
        [
            ["System Manager", "All Modules", "Full administrative control", "IT Admin, CTO"],
            ["HR Manager", "Human Resources", "Full HR + approve authority", "HR Director, HR Head"],
            ["HR User", "Human Resources", "Day-to-day HR operations", "HR Executive, HR Coordinator"],
            ["Accounts Manager", "Accounts", "Full finance + approve authority", "CFO, Finance Director"],
            ["Accounts User", "Accounts", "Day-to-day accounting", "Accountant, Bookkeeper"],
            ["Sales Manager", "Sales", "Full sales + approve authority", "Sales Director, VP Sales"],
            ["Sales User", "Sales", "Day-to-day sales operations", "Sales Executive, BDE"],
            ["Purchase Manager", "Buying", "Full procurement + approve", "Procurement Head"],
            ["Purchase User", "Buying", "Day-to-day procurement", "Purchase Executive"],
            ["Stock Manager", "Stock", "Full inventory control", "Warehouse Manager"],
            ["Stock User", "Stock", "Day-to-day stock operations", "Store Keeper, Inventory Clerk"],
            ["Manufacturing Manager", "Manufacturing", "Full production control", "Plant Manager"],
            ["Manufacturing User", "Manufacturing", "Day-to-day production", "Production Supervisor"],
            ["Projects Manager", "Projects", "Full project management", "Project Director, PMO"],
            ["Projects User", "Projects", "Task & timesheet entry", "Project Member, Developer"],
            ["Support Team", "Support", "Issue management & SLAs", "Support Agent, Help Desk"],
            ["Website Manager", "Website", "CMS & portal management", "Content Manager, Marketing"],
            ["Employee", "Self-Service", "Leave, expense, attendance", "All employees"],
            ["Auditor", "All (Read-Only)", "Compliance & audit review", "Internal/External Auditor"],
        ]
    )

    doc.add_page_break()

    add_heading(doc, "5.2 Module Access Matrix", level=2)
    add_body(doc, "The following matrix shows access levels for each role across all major modules. Badge legend: ● Full Access | ● Read/Write | ● Read Only | ● No Access")

    matrix_headers = ["Role", "HR", "Accounts", "Sales", "Buying", "Stock", "Mfg", "Projects", "Support", "Setup"]
    matrix_rows = [
        ["System Manager",     "Full Access","Full Access","Full Access","Full Access","Full Access","Full Access","Full Access","Full Access","Full Access"],
        ["HR Manager",         "Full Access","Read Only",  "No Access", "No Access", "No Access", "No Access", "Read Only", "No Access", "Limited"],
        ["HR User",            "Read/Write","No Access",  "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "No Access"],
        ["Accounts Manager",   "Read Only", "Full Access","Read Only", "Read Only", "Read Only", "No Access", "Read Only", "No Access", "Limited"],
        ["Accounts User",      "No Access", "Read/Write","Read Only", "Read Only", "No Access", "No Access", "No Access", "No Access", "No Access"],
        ["Sales Manager",      "No Access", "Read Only", "Full Access","No Access", "Read Only", "No Access", "Read Only", "Read Only", "No Access"],
        ["Sales User",         "No Access", "No Access", "Read/Write","No Access", "Read Only", "No Access", "No Access", "No Access", "No Access"],
        ["Purchase Manager",   "No Access", "Read Only", "No Access", "Full Access","Read Only", "Read Only", "No Access", "No Access", "No Access"],
        ["Purchase User",      "No Access", "No Access", "No Access", "Read/Write","Read Only", "No Access", "No Access", "No Access", "No Access"],
        ["Stock Manager",      "No Access", "No Access", "Read Only", "Read Only", "Full Access","Read Only", "No Access", "No Access", "No Access"],
        ["Stock User",         "No Access", "No Access", "No Access", "No Access", "Read/Write","No Access", "No Access", "No Access", "No Access"],
        ["Manufacturing Mgr",  "No Access", "No Access", "No Access", "Read Only", "Read/Write","Full Access","No Access", "No Access", "No Access"],
        ["Manufacturing User", "No Access", "No Access", "No Access", "No Access", "Read Only", "Read/Write","No Access", "No Access", "No Access"],
        ["Projects Manager",   "No Access", "Read Only", "No Access", "No Access", "No Access", "No Access", "Full Access","No Access", "No Access"],
        ["Projects User",      "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "Read/Write","No Access", "No Access"],
        ["Support Team",       "No Access", "No Access", "Read Only", "No Access", "No Access", "No Access", "Read Only", "Full Access","No Access"],
        ["Website Manager",    "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "Limited"],
        ["Employee",           "Limited",   "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "No Access", "No Access"],
        ["Auditor",            "Read Only", "Read Only", "Read Only", "Read Only", "Read Only", "Read Only", "Read Only", "Read Only", "Read Only"],
    ]

    table = doc.add_table(rows=1, cols=len(matrix_headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    add_header_row(table, matrix_headers)

    for idx, row_data in enumerate(matrix_rows):
        row = table.add_row()
        bg = ALT_ROW if idx % 2 == 0 else WHITE_HEX
        for i, text in enumerate(row_data):
            if i == 0:
                set_cell_shading(row.cells[i], bg)
                set_cell_text(row.cells[i], text, bold=True, size=8)
            else:
                bc, bbg = STATUS_BADGES.get(text, (DARK, WHITE_HEX))
                set_cell_shading(row.cells[i], bbg)
                cell = row.cells[i]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(f"● {text}")
                run.font.name = "Arial"
                run.font.size = Pt(7)
                run.font.bold = True
                run.font.color.rgb = bc

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 6-24. ROLE-SPECIFIC SECTIONS
    # ═══════════════════════════════════════════════════════════════════════

    roles = [
        {
            "num": 6, "title": "System Administrator",
            "role_name": "System Manager",
            "description": "The System Administrator has complete control over the ERPNext instance. This role is responsible for initial setup, user management, module configuration, data imports, workflow definitions, and system maintenance.",
            "key_responsibilities": [
                "Create and manage user accounts, assign roles and permissions",
                "Configure company settings, chart of accounts, fiscal year, and currencies",
                "Set up email domains, notification rules, and print formats",
                "Define naming series, workflows, and approval hierarchies",
                "Manage data imports/exports and bulk operations",
                "Configure backup schedules and monitor system health",
                "Set up custom fields, custom doctypes, and scripting",
                "Manage integrations (payment gateways, email, SMS, etc.)",
            ],
            "navigation": [
                ("Setup Wizard", "Setup > Setup Wizard", "Initial company configuration"),
                ("User List", "Setup > User", "Create/edit user accounts"),
                ("Role Permissions Manager", "Setup > Role Permissions Manager", "Fine-tune role access"),
                ("System Settings", "Setup > Settings > System Settings", "Session timeout, email, security"),
                ("Email Account", "Setup > Email > Email Account", "Configure outgoing/incoming email"),
                ("Workflow", "Setup > Workflow", "Define document approval flows"),
                ("Data Import", "Setup > Data > Data Import", "Bulk import from CSV/Excel"),
                ("Error Log", "Setup > Error Log", "Monitor system errors"),
                ("Scheduled Jobs", "Setup > Scheduled Jobs", "View/manage background tasks"),
                ("Backup", "Setup > Download Backups", "Download database backups"),
            ],
            "permissions": [
                ("All Modules", "Full Access", "Create, Read, Write, Delete, Submit, Amend, Cancel"),
                ("System Settings", "Full Access", "Configure all system parameters"),
                ("User Management", "Full Access", "Create users, assign roles, reset passwords"),
                ("Data Import/Export", "Full Access", "Bulk data operations"),
                ("Customization", "Full Access", "Custom fields, doctypes, scripts"),
            ],
        },
        {
            "num": 7, "title": "HR Manager",
            "role_name": "HR Manager",
            "description": "The HR Manager has full access to all Human Resources modules with approval authority. This role oversees employee lifecycle management, payroll processing, leave policies, recruitment, and HR analytics.",
            "key_responsibilities": [
                "Approve leave applications, expense claims, and travel requests",
                "Manage employee onboarding and exit processes",
                "Configure payroll components, salary structures, and tax slabs",
                "Process monthly payroll and generate payslips",
                "Define leave policies, holiday lists, and attendance rules",
                "Manage recruitment pipeline (job openings, applications, interviews)",
                "Generate HR reports and analytics dashboards",
                "Configure appraisal cycles and performance management",
            ],
            "navigation": [
                ("Employee", "HR > Employee", "View/edit employee master records"),
                ("Leave Application", "HR > Leave Application", "Approve/reject leave requests"),
                ("Payroll Entry", "HR > Payroll Entry", "Process monthly payroll"),
                ("Salary Structure", "HR > Salary Structure", "Define earning/deduction components"),
                ("Attendance", "HR > Attendance", "View/manage daily attendance"),
                ("Expense Claim", "HR > Expense Claim", "Approve employee expense claims"),
                ("Job Opening", "HR > Recruitment > Job Opening", "Post new job openings"),
                ("Appraisal", "HR > Appraisal", "Manage performance reviews"),
                ("Leave Policy", "HR > Leave Policy", "Define leave allocation rules"),
                ("HR Settings", "HR > Settings > HR Settings", "Configure HR module parameters"),
            ],
            "permissions": [
                ("Employee", "Full Access", "Create, edit, view all employee records"),
                ("Leave Application", "Approve", "Approve/reject leave requests"),
                ("Payroll Entry", "Submit", "Process and submit payroll"),
                ("Expense Claim", "Approve", "Approve/reject expense claims"),
                ("Salary Slip", "Full Access", "Generate and distribute payslips"),
                ("Recruitment", "Full Access", "Manage entire hiring pipeline"),
            ],
        },
        {
            "num": 8, "title": "HR User",
            "role_name": "HR User",
            "description": "The HR User handles day-to-day HR operations including employee data entry, attendance tracking, leave management, and recruitment coordination. This role can create documents but requires manager approval for submission.",
            "key_responsibilities": [
                "Enter and maintain employee master data",
                "Track daily attendance and manage attendance corrections",
                "Process leave applications and forward for approval",
                "Coordinate recruitment activities (schedule interviews, update status)",
                "Prepare payroll input data and assist with payroll processing",
                "Maintain training records and employee documents",
            ],
            "navigation": [
                ("Employee", "HR > Employee", "Create/edit employee records"),
                ("Attendance", "HR > Attendance", "Mark daily attendance"),
                ("Leave Application", "HR > Leave Application", "Create leave requests"),
                ("Job Applicant", "HR > Recruitment > Job Applicant", "Track applications"),
                ("Employee Checkin", "HR > Employee Checkin", "View biometric/manual checkins"),
                ("Training Event", "HR > Training > Training Event", "Schedule training sessions"),
            ],
            "permissions": [
                ("Employee", "Create/Edit", "Create and edit employee records"),
                ("Attendance", "Create/Edit", "Mark and correct attendance"),
                ("Leave Application", "Create/Edit", "Create leave requests (approval by HR Manager)"),
                ("Job Applicant", "Create/Edit", "Manage job applications"),
                ("Payroll Entry", "View Only", "View payroll data, cannot submit"),
            ],
        },
        {
            "num": 9, "title": "Accounts Manager",
            "role_name": "Accounts Manager",
            "description": "The Accounts Manager has full control over financial operations including general ledger management, accounts payable/receivable, tax configuration, budget management, and financial reporting. This role has approval authority for all financial transactions.",
            "key_responsibilities": [
                "Approve payment entries, journal entries, and financial documents",
                "Manage chart of accounts and cost centers",
                "Configure tax templates, withholding tax, and GST settings",
                "Process bank reconciliation and period closing",
                "Generate financial statements (P&L, Balance Sheet, Cash Flow)",
                "Manage budgets and budget variance analysis",
                "Handle inter-company transactions and consolidation",
                "Configure payment terms and credit limits for customers",
            ],
            "navigation": [
                ("Chart of Accounts", "Accounts > Chart of Accounts", "View/edit account hierarchy"),
                ("Journal Entry", "Accounts > Journal Entry", "Create/approve journal entries"),
                ("Payment Entry", "Accounts > Payment Entry", "Process payments"),
                ("Sales Invoice", "Accounts > Sales Invoice", "View/approve sales invoices"),
                ("Purchase Invoice", "Accounts > Purchase Invoice", "View/approve purchase invoices"),
                ("Bank Reconciliation", "Accounts > Bank Reconciliation", "Reconcile bank statements"),
                ("Trial Balance", "Accounts > Trial Balance", "View trial balance report"),
                ("P&L Statement", "Accounts > Profit and Loss", "View income/expense report"),
                ("Balance Sheet", "Accounts > Balance Sheet", "View assets/liabilities"),
                ("Budget", "Accounts > Budget", "Create and manage budgets"),
            ],
            "permissions": [
                ("Journal Entry", "Full Access", "Create, submit, amend journal entries"),
                ("Payment Entry", "Approve", "Approve/reject payments"),
                ("Sales Invoice", "Full Access", "Create, submit, cancel invoices"),
                ("Purchase Invoice", "Full Access", "Create, submit, cancel invoices"),
                ("Period Closing", "Submit", "Close accounting periods"),
                ("Budget", "Full Access", "Create and manage budgets"),
            ],
        },
        {
            "num": 10, "title": "Accounts User",
            "role_name": "Accounts User",
            "description": "The Accounts User handles routine accounting tasks including data entry for invoices, payment processing, expense recording, and basic financial reporting. Documents requiring approval are forwarded to the Accounts Manager.",
            "key_responsibilities": [
                "Enter sales and purchase invoices",
                "Process payment entries and record receipts",
                "Create journal entries for routine transactions",
                "Assist with bank reconciliation",
                "Generate day-to-day accounting reports",
                "Manage accounts receivable and payable aging",
            ],
            "navigation": [
                ("Sales Invoice", "Accounts > Sales Invoice", "Create sales invoices"),
                ("Purchase Invoice", "Accounts > Purchase Invoice", "Create purchase invoices"),
                ("Payment Entry", "Accounts > Payment Entry", "Record payments"),
                ("Journal Entry", "Accounts > Journal Entry", "Create journal entries"),
                ("Accounts Receivable", "Accounts > Accounts Receivable", "AR aging report"),
                ("Accounts Payable", "Accounts > Accounts Payable", "AP aging report"),
            ],
            "permissions": [
                ("Sales Invoice", "Create/Edit", "Create invoices (submit requires manager)"),
                ("Purchase Invoice", "Create/Edit", "Create invoices (submit requires manager)"),
                ("Payment Entry", "Create/Edit", "Record payments (approval required)"),
                ("Journal Entry", "Create/Edit", "Create entries (submit requires manager)"),
                ("Financial Reports", "View Only", "Generate and view reports"),
            ],
        },
        {
            "num": 11, "title": "Sales Manager",
            "role_name": "Sales Manager",
            "description": "The Sales Manager oversees the entire sales pipeline from lead generation to revenue collection. This role has approval authority for quotations, sales orders, and discount approvals. Access to CRM and sales analytics dashboards.",
            "key_responsibilities": [
                "Manage and assign leads to the sales team",
                "Approve quotations, discounts, and sales orders",
                "Monitor sales pipeline and conversion metrics",
                "Set and track sales targets by territory and team",
                "Approve delivery schedules and customer credit limits",
                "Generate sales reports and revenue forecasts",
                "Manage customer relationships and pricing rules",
            ],
            "navigation": [
                ("Lead", "CRM > Lead", "View/assign leads"),
                ("Opportunity", "CRM > Opportunity", "Track sales opportunities"),
                ("Quotation", "Selling > Quotation", "Create/approve quotations"),
                ("Sales Order", "Selling > Sales Order", "Approve sales orders"),
                ("Sales Invoice", "Accounts > Sales Invoice", "View sales invoices"),
                ("Delivery Note", "Stock > Delivery Note", "Track deliveries"),
                ("Sales Analytics", "Selling > Sales Analytics", "Revenue dashboards"),
                ("Pricing Rule", "Selling > Pricing Rule", "Configure pricing rules"),
            ],
            "permissions": [
                ("Lead", "Full Access", "Create, assign, convert leads"),
                ("Quotation", "Approve", "Approve quotations and discounts"),
                ("Sales Order", "Approve", "Approve sales orders"),
                ("Sales Invoice", "View Only", "View invoices (Accounts manages)"),
                ("Delivery Note", "Read Only", "Monitor delivery status"),
                ("Customer", "Full Access", "Manage customer master data"),
            ],
        },
        {
            "num": 12, "title": "Sales User",
            "role_name": "Sales User",
            "description": "The Sales User handles daily sales activities including lead follow-ups, quotation creation, order processing, and customer communication. Documents are routed to the Sales Manager for approval.",
            "key_responsibilities": [
                "Follow up on assigned leads and update status",
                "Create quotations and send to customers",
                "Convert approved quotations to sales orders",
                "Coordinate with stock team for delivery schedules",
                "Update CRM with customer interactions and notes",
            ],
            "navigation": [
                ("Lead", "CRM > Lead", "Follow up on assigned leads"),
                ("Quotation", "Selling > Quotation", "Create quotations"),
                ("Sales Order", "Selling > Sales Order", "Create sales orders"),
                ("Customer", "Selling > Customer", "View customer records"),
                ("Item", "Stock > Item", "View product catalog"),
            ],
            "permissions": [
                ("Lead", "Create/Edit", "Create and update leads"),
                ("Quotation", "Create/Edit", "Create quotations (approval required)"),
                ("Sales Order", "Create/Edit", "Create orders (approval required)"),
                ("Customer", "Read/Write", "View and update customer data"),
                ("Sales Invoice", "No Access", "Managed by Accounts team"),
            ],
        },
        {
            "num": 13, "title": "Purchase Manager",
            "role_name": "Purchase Manager",
            "description": "The Purchase Manager controls the entire procure-to-pay cycle with approval authority for purchase orders, supplier evaluations, and procurement budgets.",
            "key_responsibilities": [
                "Approve purchase orders, RFQs, and supplier quotations",
                "Manage supplier master data and supplier scorecards",
                "Define procurement policies and approval thresholds",
                "Monitor purchase analytics and cost savings",
                "Approve material requests from departments",
                "Negotiate contracts and payment terms with suppliers",
            ],
            "navigation": [
                ("Material Request", "Stock > Material Request", "Review/approve requests"),
                ("Request for Quotation", "Buying > Request for Quotation", "Send RFQs to suppliers"),
                ("Supplier Quotation", "Buying > Supplier Quotation", "Compare supplier quotes"),
                ("Purchase Order", "Buying > Purchase Order", "Approve purchase orders"),
                ("Purchase Receipt", "Stock > Purchase Receipt", "Confirm goods receipt"),
                ("Purchase Invoice", "Accounts > Purchase Invoice", "View purchase invoices"),
                ("Supplier", "Buying > Supplier", "Manage supplier records"),
                ("Purchase Analytics", "Buying > Purchase Analytics", "Procurement dashboards"),
            ],
            "permissions": [
                ("Purchase Order", "Approve", "Approve/reject purchase orders"),
                ("Material Request", "Approve", "Approve material requests"),
                ("Supplier Quotation", "Full Access", "Manage supplier comparisons"),
                ("Supplier", "Full Access", "Create and manage suppliers"),
                ("Purchase Receipt", "Read/Write", "Confirm goods received"),
            ],
        },
        {
            "num": 14, "title": "Purchase User",
            "role_name": "Purchase User",
            "description": "The Purchase User handles routine procurement tasks including creating purchase orders, managing RFQs, and coordinating with suppliers for deliveries.",
            "key_responsibilities": [
                "Create purchase orders from approved material requests",
                "Send RFQs to suppliers and collect quotations",
                "Track purchase order delivery status",
                "Coordinate with warehouse for goods receipt",
                "Maintain supplier communication records",
            ],
            "navigation": [
                ("Material Request", "Stock > Material Request", "Create material requests"),
                ("Request for Quotation", "Buying > Request for Quotation", "Create and send RFQs"),
                ("Purchase Order", "Buying > Purchase Order", "Create purchase orders"),
                ("Purchase Receipt", "Stock > Purchase Receipt", "Record goods receipt"),
                ("Supplier", "Buying > Supplier", "View supplier records"),
            ],
            "permissions": [
                ("Purchase Order", "Create/Edit", "Create POs (approval required)"),
                ("Material Request", "Create/Edit", "Create requests"),
                ("RFQ", "Create/Edit", "Send RFQs to suppliers"),
                ("Purchase Receipt", "Create/Edit", "Record goods received"),
                ("Purchase Invoice", "No Access", "Managed by Accounts team"),
            ],
        },
        {
            "num": 15, "title": "Stock Manager",
            "role_name": "Stock Manager",
            "description": "The Stock Manager has full control over inventory operations including warehouse management, stock transfers, serial/batch tracking, and inventory valuation.",
            "key_responsibilities": [
                "Manage warehouse setup and bin locations",
                "Approve stock transfers and material issues",
                "Conduct stock reconciliation and cycle counts",
                "Configure serial number and batch tracking policies",
                "Monitor reorder levels and generate purchase requests",
                "Manage inventory valuation and costing methods",
            ],
            "navigation": [
                ("Stock Entry", "Stock > Stock Entry", "Material receipt/issue/transfer"),
                ("Stock Reconciliation", "Stock > Stock Reconciliation", "Adjust stock levels"),
                ("Delivery Note", "Stock > Delivery Note", "Ship goods to customers"),
                ("Purchase Receipt", "Stock > Purchase Receipt", "Receive goods from suppliers"),
                ("Warehouse", "Stock > Warehouse", "Manage warehouse hierarchy"),
                ("Stock Balance", "Stock > Stock Balance", "View current stock levels"),
                ("Stock Ledger", "Stock > Stock Ledger", "View stock transactions"),
                ("Item", "Stock > Item", "Manage item master data"),
            ],
            "permissions": [
                ("Stock Entry", "Full Access", "All stock movements"),
                ("Stock Reconciliation", "Submit", "Submit stock adjustments"),
                ("Warehouse", "Full Access", "Create and manage warehouses"),
                ("Item", "Full Access", "Manage item master"),
                ("Delivery Note", "Full Access", "Ship and track deliveries"),
            ],
        },
        {
            "num": 16, "title": "Stock User",
            "role_name": "Stock User",
            "description": "The Stock User handles day-to-day warehouse operations including receiving goods, issuing materials, and updating stock records.",
            "key_responsibilities": [
                "Record goods receipt against purchase orders",
                "Issue materials to production or departments",
                "Process stock transfers between warehouses",
                "Update stock records and maintain accuracy",
                "Report stock discrepancies to Stock Manager",
            ],
            "navigation": [
                ("Stock Entry", "Stock > Stock Entry", "Create stock movements"),
                ("Purchase Receipt", "Stock > Purchase Receipt", "Receive goods"),
                ("Delivery Note", "Stock > Delivery Note", "Dispatch goods"),
                ("Stock Balance", "Stock > Stock Balance", "Check stock levels"),
                ("Item", "Stock > Item", "View item details"),
            ],
            "permissions": [
                ("Stock Entry", "Create/Edit", "Create entries (submit requires manager)"),
                ("Purchase Receipt", "Create/Edit", "Record receipts"),
                ("Delivery Note", "Create/Edit", "Create delivery notes"),
                ("Item", "Read Only", "View item master (cannot edit)"),
                ("Stock Reconciliation", "No Access", "Manager only"),
            ],
        },
        {
            "num": 17, "title": "Manufacturing Manager",
            "role_name": "Manufacturing Manager",
            "description": "The Manufacturing Manager controls production planning, BOM management, work order execution, and quality control processes.",
            "key_responsibilities": [
                "Create and manage Bills of Materials (BOMs)",
                "Plan production schedules and capacity",
                "Approve and monitor work orders",
                "Manage quality inspection criteria and results",
                "Track production costs and efficiency metrics",
                "Coordinate with procurement for raw materials",
            ],
            "navigation": [
                ("BOM", "Manufacturing > BOM", "Create/edit Bills of Materials"),
                ("Production Plan", "Manufacturing > Production Plan", "Plan production batches"),
                ("Work Order", "Manufacturing > Work Order", "Create/monitor work orders"),
                ("Job Card", "Manufacturing > Job Card", "Track work station operations"),
                ("Quality Inspection", "Stock > Quality Inspection", "Define/review inspections"),
                ("BOM Stock Report", "Manufacturing > BOM Stock Report", "Check material availability"),
            ],
            "permissions": [
                ("BOM", "Full Access", "Create, edit, submit BOMs"),
                ("Work Order", "Approve", "Approve and submit work orders"),
                ("Production Plan", "Full Access", "Create and execute production plans"),
                ("Quality Inspection", "Full Access", "Define criteria and review results"),
                ("Job Card", "Full Access", "Manage work station jobs"),
            ],
        },
        {
            "num": 18, "title": "Manufacturing User",
            "role_name": "Manufacturing User",
            "description": "The Manufacturing User executes production operations including work order processing, job card updates, and quality inspection recording.",
            "key_responsibilities": [
                "Execute assigned work orders and update progress",
                "Record production quantities and scrap",
                "Complete job cards at each work station",
                "Record quality inspection results",
                "Report production issues to Manufacturing Manager",
            ],
            "navigation": [
                ("Work Order", "Manufacturing > Work Order", "View and update work orders"),
                ("Job Card", "Manufacturing > Job Card", "Complete job card entries"),
                ("Quality Inspection", "Stock > Quality Inspection", "Record inspection results"),
                ("Stock Entry", "Stock > Stock Entry", "Material transfer for manufacturing"),
                ("BOM", "Manufacturing > BOM", "View BOMs (read only)"),
            ],
            "permissions": [
                ("Work Order", "Read/Write", "Update progress (cannot approve)"),
                ("Job Card", "Create/Edit", "Create and complete job cards"),
                ("Quality Inspection", "Create/Edit", "Record inspection results"),
                ("BOM", "View Only", "View BOMs, cannot edit"),
                ("Production Plan", "View Only", "View plans, cannot modify"),
            ],
        },
        {
            "num": 19, "title": "Projects Manager",
            "role_name": "Projects Manager",
            "description": "The Projects Manager has full control over project creation, task assignment, resource allocation, timesheet approval, and project profitability analysis.",
            "key_responsibilities": [
                "Create and configure projects with milestones",
                "Assign tasks to team members with deadlines",
                "Approve timesheets and track billable hours",
                "Monitor project progress via Gantt charts and dashboards",
                "Analyze project profitability and resource utilization",
                "Manage project templates for recurring project types",
            ],
            "navigation": [
                ("Project", "Projects > Project", "Create/manage projects"),
                ("Task", "Projects > Task", "Assign and track tasks"),
                ("Timesheet", "Projects > Timesheet", "Approve team timesheets"),
                ("Activity Type", "Projects > Activity Type", "Define activity categories"),
                ("Project Gantt", "Projects > Project > Gantt", "Visual project timeline"),
                ("Daily Timesheet Summary", "Projects > Daily Timesheet Summary", "Team time report"),
            ],
            "permissions": [
                ("Project", "Full Access", "Create, edit, close projects"),
                ("Task", "Full Access", "Create, assign, manage tasks"),
                ("Timesheet", "Approve", "Approve/reject timesheets"),
                ("Activity Type", "Full Access", "Define activity categories"),
                ("Project Reports", "Full Access", "All project analytics"),
            ],
        },
        {
            "num": 20, "title": "Projects User",
            "role_name": "Projects User",
            "description": "The Projects User participates in projects by completing assigned tasks, logging timesheets, and updating task progress.",
            "key_responsibilities": [
                "View and complete assigned tasks",
                "Log daily/weekly timesheets with activity details",
                "Update task status and progress percentage",
                "Add comments and attachments to tasks",
                "View project Gantt charts and milestones",
            ],
            "navigation": [
                ("Task", "Projects > Task", "View and update assigned tasks"),
                ("Timesheet", "Projects > Timesheet", "Log work hours"),
                ("Project", "Projects > Project", "View project overview"),
                ("Activity Type", "Projects > Activity Type", "Select activity for timesheet"),
            ],
            "permissions": [
                ("Task", "Read/Write", "Update assigned tasks only"),
                ("Timesheet", "Create/Edit", "Log timesheets (approval required)"),
                ("Project", "View Only", "View project details"),
                ("Activity Type", "Read Only", "View activity categories"),
            ],
        },
        {
            "num": 21, "title": "Support Team",
            "role_name": "Support Team",
            "description": "The Support Team manages customer issues, service requests, warranty claims, and knowledge base content. This role handles the complete support lifecycle from ticket creation to resolution.",
            "key_responsibilities": [
                "Create and manage customer support issues/tickets",
                "Assign issues to team members based on expertise",
                "Track SLA compliance and response times",
                "Manage warranty claims and service contracts",
                "Maintain knowledge base articles for self-service",
                "Generate support metrics and customer satisfaction reports",
            ],
            "navigation": [
                ("Issue", "Support > Issue", "Create/manage support tickets"),
                ("Service Level Agreement", "Support > SLA", "Define/monitor SLAs"),
                ("Warranty Claim", "Support > Warranty Claim", "Process warranty requests"),
                ("Issue Type", "Support > Issue Type", "Categorize issue types"),
                ("Issue Priority", "Support > Issue Priority", "Set priority levels"),
                ("Knowledge Base", "Support > Knowledge Base", "Manage help articles"),
            ],
            "permissions": [
                ("Issue", "Full Access", "Create, assign, resolve, close issues"),
                ("SLA", "Read/Write", "View and update SLA records"),
                ("Warranty Claim", "Create/Edit", "Process warranty claims"),
                ("Knowledge Base", "Full Access", "Create and edit articles"),
                ("Customer", "Read Only", "View customer details"),
            ],
        },
        {
            "num": 22, "title": "Website Manager",
            "role_name": "Website Manager",
            "description": "The Website Manager controls the public-facing website built into ERPNext, including content management, blog posts, web forms, and portal configuration.",
            "key_responsibilities": [
                "Create and manage web pages and landing pages",
                "Publish blog posts and news articles",
                "Configure web forms for lead capture and feedback",
                "Manage website settings (logo, favicon, navigation)",
                "Set up customer/supplier portal access",
                "Monitor website analytics and page performance",
            ],
            "navigation": [
                ("Web Page", "Website > Web Page", "Create/edit web pages"),
                ("Blog Post", "Website > Blog Post", "Publish blog articles"),
                ("Web Form", "Website > Web Form", "Create online forms"),
                ("Website Settings", "Website > Website Settings", "Configure site settings"),
                ("Website Theme", "Website > Website Theme", "Customize appearance"),
                ("Navbar/Footer", "Website > Top Bar / Footer", "Configure navigation"),
            ],
            "permissions": [
                ("Web Page", "Full Access", "Create, edit, publish pages"),
                ("Blog Post", "Full Access", "Create and publish blogs"),
                ("Web Form", "Full Access", "Create and manage forms"),
                ("Website Settings", "Full Access", "Configure site settings"),
                ("Portal Settings", "Full Access", "Configure portal access"),
            ],
        },
        {
            "num": 23, "title": "Employee (Self-Service)",
            "role_name": "Employee",
            "description": "The Employee role provides self-service access for all staff members. Employees can apply for leave, submit expense claims, mark attendance, view payslips, and update personal information without requiring HR intervention.",
            "key_responsibilities": [
                "Apply for leave and track leave balance",
                "Submit expense claims with supporting documents",
                "View and download monthly payslips",
                "Update personal information (address, emergency contact, bank details)",
                "View attendance records and regularize discrepancies",
                "Complete assigned training programs",
                "Participate in performance appraisals",
            ],
            "navigation": [
                ("Leave Application", "HR > Leave Application", "Apply for leave"),
                ("Expense Claim", "HR > Expense Claim", "Submit expense claims"),
                ("Salary Slip", "HR > Salary Slip", "View/download payslips"),
                ("Attendance", "HR > Attendance", "View attendance records"),
                ("Employee Profile", "HR > Employee", "Update personal details"),
                ("Leave Balance", "HR > Leave Balance", "Check remaining leave"),
                ("Tax Declaration", "HR > Employee Tax Exemption Declaration", "Submit tax savings"),
            ],
            "permissions": [
                ("Leave Application", "Create/Edit", "Apply for own leave only"),
                ("Expense Claim", "Create/Edit", "Submit own expenses only"),
                ("Salary Slip", "View Only", "View own payslips only"),
                ("Attendance", "View Only", "View own attendance only"),
                ("Employee", "Limited", "Edit own profile fields only"),
                ("Tax Declaration", "Create/Edit", "Submit own tax declarations"),
            ],
        },
        {
            "num": 24, "title": "Auditor (Read-Only)",
            "role_name": "Auditor",
            "description": "The Auditor role provides comprehensive read-only access across all modules for compliance reviews, internal audits, and financial verification. Auditors cannot create, edit, or delete any records.",
            "key_responsibilities": [
                "Review financial transactions and journal entries",
                "Verify compliance with accounting standards",
                "Audit HR records and payroll accuracy",
                "Review procurement processes and vendor compliance",
                "Examine inventory records and stock valuation",
                "Generate audit reports and flag discrepancies",
                "Verify workflow adherence and approval chains",
            ],
            "navigation": [
                ("General Ledger", "Accounts > General Ledger", "Review all transactions"),
                ("Audit Trail", "Setup > Audit Trail", "Track document changes"),
                ("Activity Log", "Setup > Activity Log", "View user activity"),
                ("Sales Invoice", "Accounts > Sales Invoice", "Review invoices"),
                ("Purchase Invoice", "Accounts > Purchase Invoice", "Review invoices"),
                ("Stock Ledger", "Stock > Stock Ledger", "Review stock movements"),
                ("Employee", "HR > Employee", "Review employee records"),
                ("Workflow Actions", "Setup > Workflow Actions", "Review approval flows"),
            ],
            "permissions": [
                ("All Modules", "Read Only", "View all records across all modules"),
                ("Reports", "Read Only", "Generate and view all reports"),
                ("Audit Trail", "Read Only", "Review document modification history"),
                ("Activity Log", "Read Only", "Monitor user activities"),
                ("Create/Edit/Delete", "No Access", "Cannot modify any data"),
            ],
        },
    ]

    for role in roles:
        add_heading(doc, f"{role['num']}. {role['title']}", level=1)

        # Role badge
        p = doc.add_paragraph()
        run = p.add_run(f"ERPNext Role: ")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = DARK
        run = p.add_run(role['role_name'])
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = BLUE

        add_body(doc, role['description'])

        add_heading(doc, f"{role['num']}.1 Key Responsibilities", level=2)
        for resp in role['key_responsibilities']:
            add_bullet(doc, resp)

        add_heading(doc, f"{role['num']}.2 Navigation Guide", level=2)
        nav_table = make_table(doc,
            ["Feature", "Navigation Path", "Description"],
            [[n[0], n[1], n[2]] for n in role['navigation']]
        )

        add_heading(doc, f"{role['num']}.3 Permissions", level=2)
        make_table(doc,
            ["Document/Module", "Access Level", "Details"],
            [[p[0], p[1], p[2]] for p in role['permissions']],
            badge_col=1,
            badge_map=STATUS_BADGES
        )

        doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 25. DASHBOARD & REPORTS
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "25. Dashboard & Reports Guide", level=1)

    add_body(doc, "ERPNext provides built-in dashboards and reports for every module. Each role has access to reports relevant to their function. Below is a guide to the most commonly used reports.")

    add_heading(doc, "25.1 Accessing Reports", level=2)
    add_bullet(doc, "From the sidebar: Navigate to the relevant module, then click 'Reports' in the sidebar")
    add_bullet(doc, "From the search bar: Press Ctrl+K (or Cmd+K on Mac) and type the report name")
    add_bullet(doc, "From the desk: Click on report cards visible on module home pages")

    add_heading(doc, "25.2 Key Reports by Module", level=2)

    make_table(doc,
        ["Module", "Report Name", "Description", "Available To"],
        [
            ["Accounts", "General Ledger", "All transactions for an account/period", "Accounts Manager, Auditor"],
            ["Accounts", "Profit and Loss", "Income vs expenses summary", "Accounts Manager, Auditor"],
            ["Accounts", "Balance Sheet", "Assets, liabilities, equity snapshot", "Accounts Manager, Auditor"],
            ["Accounts", "Accounts Receivable", "Customer outstanding aging", "Accounts Manager/User"],
            ["Accounts", "Accounts Payable", "Supplier outstanding aging", "Accounts Manager/User"],
            ["Sales", "Sales Analytics", "Revenue by customer/item/territory", "Sales Manager"],
            ["Sales", "Sales Funnel", "Lead to order conversion pipeline", "Sales Manager"],
            ["Buying", "Purchase Analytics", "Spend by supplier/item/category", "Purchase Manager"],
            ["Stock", "Stock Balance", "Current inventory by warehouse", "Stock Manager/User"],
            ["Stock", "Stock Ledger", "All stock movements chronologically", "Stock Manager"],
            ["HR", "Employee Leave Balance", "Leave allocation vs used", "HR Manager/User"],
            ["HR", "Monthly Attendance Sheet", "Team attendance for a month", "HR Manager/User"],
            ["Manufacturing", "Production Analytics", "Work order completion rates", "Manufacturing Manager"],
            ["Projects", "Daily Timesheet Summary", "Team hours logged per day", "Projects Manager"],
        ]
    )

    add_heading(doc, "25.3 Custom Reports", level=2)
    add_body(doc, "Users with appropriate permissions can create custom reports using the Report Builder (Ctrl+K > 'Report Builder'). Custom reports can be saved, shared with specific roles, and added to dashboards.")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 26. COMMON WORKFLOWS
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "26. Common Workflows", level=1)

    add_heading(doc, "26.1 Procure-to-Pay (P2P)", level=2)
    steps = [
        ("1", "Material Request", "Department raises a material request", "Purchase User / Any"),
        ("2", "Supplier Quotation", "RFQ sent, quotes collected", "Purchase User"),
        ("3", "Purchase Order", "PO created and approved", "Purchase User / Manager"),
        ("4", "Purchase Receipt", "Goods received at warehouse", "Stock User"),
        ("5", "Purchase Invoice", "Supplier invoice recorded", "Accounts User"),
        ("6", "Payment Entry", "Payment made to supplier", "Accounts User / Manager"),
    ]
    make_table(doc, ["Step", "Document", "Action", "Role"], steps)

    add_heading(doc, "26.2 Order-to-Cash (O2C)", level=2)
    steps = [
        ("1", "Lead / Opportunity", "New prospect identified", "Sales User"),
        ("2", "Quotation", "Price quote sent to customer", "Sales User"),
        ("3", "Sales Order", "Customer confirms order", "Sales User / Manager"),
        ("4", "Delivery Note", "Goods shipped to customer", "Stock User"),
        ("5", "Sales Invoice", "Invoice sent to customer", "Accounts User"),
        ("6", "Payment Entry", "Payment received from customer", "Accounts User"),
    ]
    make_table(doc, ["Step", "Document", "Action", "Role"], steps)

    add_heading(doc, "26.3 Employee Leave Request", level=2)
    steps = [
        ("1", "Leave Application", "Employee applies for leave", "Employee"),
        ("2", "Approval", "Leave approver reviews and approves/rejects", "HR Manager / Leave Approver"),
        ("3", "Leave Allocation Update", "System deducts from leave balance", "Automatic"),
        ("4", "Attendance", "Leave dates marked in attendance", "Automatic"),
    ]
    make_table(doc, ["Step", "Document", "Action", "Role"], steps)

    add_heading(doc, "26.4 Expense Claim", level=2)
    steps = [
        ("1", "Expense Claim", "Employee submits claim with receipts", "Employee"),
        ("2", "Approval", "Manager reviews and approves", "HR Manager / Expense Approver"),
        ("3", "Journal Entry", "Accounting entry created", "Automatic"),
        ("4", "Payment Entry", "Reimbursement processed", "Accounts User"),
    ]
    make_table(doc, ["Step", "Document", "Action", "Role"], steps)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 27. SECURITY & BEST PRACTICES
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "27. Security & Best Practices", level=1)

    add_heading(doc, "27.1 Account Security", level=2)
    add_bullet(doc, "Never share your login credentials with anyone")
    add_bullet(doc, "Always log out when leaving your workstation (click avatar > Log Out)")
    add_bullet(doc, "Use a strong, unique password and change it every 90 days")
    add_bullet(doc, "Report any suspicious activity to the System Administrator immediately")
    add_bullet(doc, "Do not save passwords in browsers on shared computers")

    add_heading(doc, "27.2 Data Handling", level=2)
    add_bullet(doc, "Do not export sensitive data (payroll, financials) without authorization")
    add_bullet(doc, "Verify document details before submitting (submitted documents cannot be easily changed)")
    add_bullet(doc, "Use the 'Amend' feature instead of creating duplicate documents when corrections are needed")
    add_bullet(doc, "Always add comments when making significant changes to records")
    add_bullet(doc, "Use file attachments for supporting documents (invoices, receipts, approvals)")

    add_heading(doc, "27.3 Session Management", level=2)
    add_bullet(doc, "Sessions automatically expire after 6 hours of inactivity (configurable)")
    add_bullet(doc, "Multiple simultaneous sessions from different devices are allowed but monitored")
    add_bullet(doc, "The System Administrator can view active sessions and force logout if needed")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 28. TROUBLESHOOTING & FAQ
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "28. Troubleshooting & FAQ", level=1)

    faq = [
        ("I forgot my password", "Click 'Forgot Password' on the login page. A reset link will be sent to your registered email. If you don't receive it, contact the System Administrator."),
        ("I can't see a module/menu item", "Your role may not have access to that module. Contact your manager or the System Administrator to request the appropriate role."),
        ("I submitted a document with errors", "Use the 'Amend' button to create a corrected version. Submitted documents are locked to maintain audit trails. If Amend is not available, contact the Accounts Manager or System Administrator."),
        ("The page is loading slowly", "Clear your browser cache (Ctrl+Shift+Delete). Try using Chrome or Firefox for best performance. If the issue persists, report it to the System Administrator."),
        ("I see 'Not Permitted' error", "You don't have the required permission for this action. Check with your manager or System Administrator for role assignment."),
        ("My report shows no data", "Check the date filters and ensure you have the correct company selected. Some reports require specific filters to display data."),
        ("Email notifications are not working", "Contact the System Administrator to verify email configuration. Check your spam/junk folder as well."),
        ("I need to delete a submitted document", "Submitted documents cannot be deleted directly. They must be cancelled first (if your role permits), or contact the Accounts Manager/System Administrator."),
        ("How do I change my display language?", "Go to your User Settings (click avatar > My Settings) and change the Language field. Refresh the page for changes to take effect."),
        ("Where are my saved reports?", "Navigate to the module, click 'Report', and look for 'Saved Reports' dropdown. Custom reports you created will be listed there."),
    ]

    for q, a in faq:
        p = doc.add_paragraph()
        run = p.add_run(f"Q: {q}")
        run.font.name = "Arial"
        run.font.size = Pt(10)
        run.font.bold = True
        run.font.color.rgb = NAVY
        p2 = doc.add_paragraph()
        run2 = p2.add_run(f"A: {a}")
        run2.font.name = "Arial"
        run2.font.size = Pt(10)
        run2.font.color.rgb = DARK
        p2.paragraph_format.space_after = Pt(8)

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 29. SUPPORT & ESCALATION
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "29. Support & Escalation", level=1)

    add_body(doc, "For issues that cannot be resolved through self-service, use the following escalation matrix:")

    make_table(doc,
        ["Level", "Contact", "Scope", "Response Time"],
        [
            ["L1 — Self Service", "This User Guide + FAQ", "Common questions and how-to", "Immediate"],
            ["L2 — IT Help Desk", "System Administrator", "Account access, permissions, configuration", "< 4 hours"],
            ["L3 — Module Expert", "Department Head / Module Owner", "Business logic, workflow, process changes", "< 8 hours"],
            ["L4 — Vendor Support", "ERPNext / Frappe Support", "System bugs, performance, infrastructure", "< 24 hours"],
        ]
    )

    add_heading(doc, "29.1 Reporting Issues", level=2)
    add_body(doc, "When reporting an issue, include the following information:")
    add_bullet(doc, "Your username and role")
    add_bullet(doc, "The exact URL or page where the issue occurred")
    add_bullet(doc, "Steps to reproduce the issue")
    add_bullet(doc, "Screenshots or screen recordings if possible")
    add_bullet(doc, "The expected behavior vs actual behavior")
    add_bullet(doc, "The date and time the issue first occurred")

    doc.add_page_break()

    # ═══════════════════════════════════════════════════════════════════════
    # 30. GLOSSARY
    # ═══════════════════════════════════════════════════════════════════════

    add_heading(doc, "30. Glossary", level=1)

    glossary = [
        ("BOM", "Bill of Materials — a list of raw materials and sub-assemblies needed to manufacture a product"),
        ("CoA", "Chart of Accounts — the hierarchical structure of all ledger accounts"),
        ("CRM", "Customer Relationship Management — module for managing leads and opportunities"),
        ("DocType", "Document Type — the fundamental building block in ERPNext; each form/record type is a DocType"),
        ("Fiscal Year", "The 12-month accounting period used for financial reporting"),
        ("GL Entry", "General Ledger Entry — an individual debit or credit posting in the ledger"),
        ("JV", "Journal Voucher/Entry — manual accounting entry with debits and credits"),
        ("Material Request", "Internal request for procurement or transfer of materials"),
        ("Naming Series", "Auto-generated document numbering pattern (e.g., INV-2026-00001)"),
        ("PO", "Purchase Order — formal order to a supplier for goods/services"),
        ("RBAC", "Role-Based Access Control — security model where permissions are assigned to roles"),
        ("RFQ", "Request for Quotation — solicitation sent to suppliers for price quotes"),
        ("SLA", "Service Level Agreement — defines response and resolution time commitments"),
        ("SO", "Sales Order — confirmed customer order for goods/services"),
        ("Submit", "Action that locks a document and triggers downstream processes (e.g., GL entries)"),
        ("Workflow", "Automated approval chain defining document state transitions"),
    ]

    make_table(doc, ["Term", "Definition"], glossary)

    doc.add_paragraph()

    # ═══════════════════════════════════════════════════════════════════════
    # FOOTER BLOCK
    # ═══════════════════════════════════════════════════════════════════════

    doc.add_paragraph()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Prepared By")
    run.font.name = "Arial"
    run.font.size = Pt(12)
    run.font.bold = True
    run.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{AUTHOR}\n{COMPANY}")
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(CLASSIFICATION)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.bold = True
    run.font.color.rgb = RED

    # ── Save ─────────────────────────────────────────────────────────────

    output_path = "/Users/tpe/VSJWORK/erpnext-vsjailabs/docs/ERPNext_End_User_Guide_v1.0.docx"
    doc.save(output_path)
    print(f"Document saved: {output_path}")
    print(f"Sections: 30")
    print(f"Roles covered: 19")

if __name__ == "__main__":
    generate()
