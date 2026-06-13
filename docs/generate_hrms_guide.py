#!/usr/bin/env python3
"""
Generate Frappe HR (HRMS) End-to-End Functional Guide — Enterprise Grade DOCX
For: VSJ AI Labs
Author: Ashish Kumar Satyam
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml, OxmlElement
import datetime

# ── Brand Colors ─────────────────────────────────────────────────────────────
NAVY     = RGBColor(0x1B, 0x3A, 0x5C)
BLUE     = RGBColor(0x2E, 0x75, 0xB6)
GREEN    = RGBColor(0x27, 0xAE, 0x60)
ORANGE   = RGBColor(0xE6, 0x7E, 0x22)
RED      = RGBColor(0xE7, 0x4C, 0x3C)
DARK     = RGBColor(0x2C, 0x3E, 0x50)
MEDIUM   = RGBColor(0x7F, 0x8C, 0x8D)
WHITE    = RGBColor(0xFF, 0xFF, 0xFF)
NAVY_HEX  = "1B3A5C"
WHITE_HEX = "FFFFFF"
ALT_ROW   = "F8F9FA"
GREEN_BG  = "E8F5E9"
ORANGE_BG = "FFF8E1"
RED_BG    = "FFEBEE"

TODAY = datetime.date.today().strftime("%B %d, %Y")
VERSION = "1.0"
DOC_TITLE = "Frappe HR (HRMS) — Functional Guide"
DOC_SUBTITLE = "End-to-End Setup & Operations Manual"
COMPANY = "VSJ AI Labs"
AUTHOR = "Ashish Kumar Satyam"
CLASSIFICATION = "Internal / Confidential"
PORTAL = "https://erp.vsjailabs.in"

CODE = "Consolas"

# ── Helpers ──────────────────────────────────────────────────────────────────
def shade(cell, color_hex):
    cell._tc.get_or_add_tcPr().append(
        parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}" w:val="clear"/>'))

def cell_text(cell, text, bold=False, color=None, size=9, font="Arial", align=None, code=False):
    cell.text = ""
    p = cell.paragraphs[0]
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.font.name = CODE if code else font
    run.font.size = Pt(size)
    run.font.bold = bold
    if color:
        run.font.color.rgb = color
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)

def header_row(table, texts):
    for i, t in enumerate(texts):
        shade(table.rows[0].cells[i], NAVY_HEX)
        cell_text(table.rows[0].cells[i], t, bold=True, color=WHITE, size=9)

def data_row(table, texts, idx, code_cols=()):
    row = table.add_row()
    bg = ALT_ROW if idx % 2 == 0 else WHITE_HEX
    for i, t in enumerate(texts):
        shade(row.cells[i], bg)
        cell_text(row.cells[i], str(t), size=9, code=(i in code_cols))
    return row

def make_table(doc, headers, rows, widths=None, code_cols=()):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    header_row(table, headers)
    for idx, r in enumerate(rows):
        data_row(table, r, idx, code_cols=code_cols)
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Cm(w)
    return table

def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = NAVY
        run.font.name = "Arial"
    return h

def body(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    run.font.bold = bold
    return p

def bullet(doc, text, num=False):
    p = doc.add_paragraph(style=("List Number" if num else "List Bullet"))
    p.clear()
    run = p.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)
    run.font.color.rgb = DARK
    return p

def note(doc, title, text, color=BLUE):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{title}: ")
    r1.font.name = "Arial"; r1.font.size = Pt(10); r1.font.bold = True; r1.font.color.rgb = color
    r2 = p.add_run(text)
    r2.font.name = "Arial"; r2.font.size = Pt(10); r2.font.color.rgb = DARK
    return p

def add_page_number_footer(section):
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.font.name = "Arial"; run.font.size = Pt(8); run.font.color.rgb = MEDIUM
    run.add_text(f"{COMPANY}  |  {DOC_TITLE}  |  Page ")
    fld1 = OxmlElement('w:fldSimple'); fld1.set(qn('w:instr'), 'PAGE')
    run._r.append(fld1)
    run2 = p.add_run(); run2.font.name = "Arial"; run2.font.size = Pt(8); run2.font.color.rgb = MEDIUM
    run2.add_text(" of ")
    fld2 = OxmlElement('w:fldSimple'); fld2.set(qn('w:instr'), 'NUMPAGES')
    run2._r.append(fld2)

# ── Build ────────────────────────────────────────────────────────────────────
def generate():
    doc = Document()
    s = doc.sections[0]
    s.page_width = Cm(21.59); s.page_height = Cm(27.94)
    s.top_margin = s.bottom_margin = s.left_margin = s.right_margin = Cm(2.54)

    st = doc.styles["Normal"]; st.font.name = "Arial"; st.font.size = Pt(10); st.font.color.rgb = DARK
    add_page_number_footer(s)

    # ── COVER ──
    for _ in range(6):
        doc.add_paragraph()
    for txt, sz, col, bold in [(COMPANY, 32, NAVY, True), (DOC_TITLE, 22, BLUE, False), (DOC_SUBTITLE, 14, MEDIUM, False)]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(txt); r.font.name = "Arial"; r.font.size = Pt(sz); r.font.color.rgb = col; r.font.bold = bold
    doc.add_paragraph(); doc.add_paragraph()
    meta = doc.add_table(rows=7, cols=2); meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, (k, v) in enumerate([
        ("Document Version", VERSION), ("Date", TODAY), ("Author", AUTHOR),
        ("For", COMPANY), ("Application", "Frappe HR (HRMS) v15.60.4 on ERPNext v15"),
        ("Portal URL", PORTAL), ("Classification", CLASSIFICATION)]):
        bg = ALT_ROW if i % 2 == 0 else WHITE_HEX
        shade(meta.rows[i].cells[0], bg); shade(meta.rows[i].cells[1], bg)
        cell_text(meta.rows[i].cells[0], k, bold=True, size=10)
        cell_text(meta.rows[i].cells[1], v, bold=(k == "Classification"),
                  color=(RED if k == "Classification" else None), size=10)
    doc.add_paragraph()
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(CLASSIFICATION); r.font.name = "Arial"; r.font.size = Pt(12); r.font.bold = True; r.font.color.rgb = RED
    doc.add_page_break()

    # ── TOC ──
    heading(doc, "Table of Contents", level=1)
    toc = [
        "1. Document Information", "2. Introduction & Purpose", "3. Accessing HRMS",
        "4. Roles & Permissions", "5. One-Time Setup (Prerequisites)", "6. Employee Management",
        "7. Attendance & Shift Management", "8. Leave Management", "9. Payroll",
        "10. Expense Claims", "11. Performance Management", "12. Recruitment",
        "13. Employee Lifecycle", "14. Reports & Dashboards", "15. Mobile App (Frappe HR)",
        "16. Recommended First-Run Path", "17. Email & Notifications", "18. Troubleshooting & FAQ",
        "19. Support & Escalation",
    ]
    for t in toc:
        p = doc.add_paragraph(); r = p.add_run(t)
        r.font.name = "Arial"; r.font.size = Pt(10); r.font.color.rgb = DARK
    doc.add_page_break()

    # ── 1. DOCUMENT INFORMATION ──
    heading(doc, "1. Document Information", 1)
    make_table(doc, ["Field", "Value"], [
        ["Document Name", DOC_TITLE], ["Version", VERSION], ["Date", TODAY], ["Owner", AUTHOR],
        ["Audience", "HR Managers, HR Users, all Employees (self-service)"],
        ["Application", "Frappe HR (HRMS) v15.60.4"],
        ["Platform", "ERPNext v15 on Utho Cloud (Docker)"],
        ["Portal URL", PORTAL], ["Classification", CLASSIFICATION],
    ], widths=[5, 11], code_cols=(1,) if False else ())
    doc.add_paragraph()
    make_table(doc, ["Version", "Date", "Author", "Changes"], [
        ["1.0", TODAY, AUTHOR, "Initial release — end-to-end HRMS functional guide"],
    ])
    doc.add_page_break()

    # ── 2. INTRODUCTION ──
    heading(doc, "2. Introduction & Purpose", 1)
    body(doc, "Frappe HR (HRMS) is the open-source Human Resource Management module integrated with ERPNext. It covers the complete employee lifecycle — from recruitment and onboarding through attendance, leave, payroll, performance, and exit — within the same platform used for the rest of the business.")
    body(doc, "This guide provides a practical, end-to-end walkthrough: how to access HRMS, the one-time setup required before transactions will work, and the day-to-day operation of each functional area. It is written for VSJ AI Labs' HR team and employees on the live instance at " + PORTAL + ".")
    heading(doc, "2.1 Functional Areas Covered", 2)
    make_table(doc, ["Area", "What it does", "Key Documents (DocTypes)"], [
        ["Employee", "Master records, onboarding, lifecycle", "Employee, Employee Onboarding"],
        ["Attendance", "Daily attendance, check-in/out, shifts", "Attendance, Employee Checkin, Shift Type"],
        ["Leave", "Leave types, policies, applications, balances", "Leave Type, Leave Policy, Leave Application"],
        ["Payroll", "Salary structures, monthly payroll, payslips", "Salary Structure, Payroll Entry, Salary Slip"],
        ["Expenses", "Employee expense claims & reimbursement", "Expense Claim"],
        ["Performance", "Appraisals, goals/KRAs, feedback", "Appraisal, Appraisal Cycle"],
        ["Recruitment", "Job openings to offers", "Job Opening, Job Applicant, Job Offer"],
        ["Lifecycle", "Promotion, transfer, separation", "Employee Promotion / Transfer / Separation"],
    ], widths=[3.5, 6.5, 6])
    doc.add_page_break()

    # ── 3. ACCESSING HRMS ──
    heading(doc, "3. Accessing HRMS", 1)
    body(doc, "HRMS is organised into workspaces in the ERPNext desk. After login, use the left sidebar or the direct paths below.")
    heading(doc, "3.1 Access Points", 2)
    make_table(doc, ["What", "Path / URL"], [
        ["Main desk", f"{PORTAL}/app"],
        ["HR workspace", f"{PORTAL}/app/hr"],
        ["Payroll workspace", f"{PORTAL}/app/payroll"],
        ["Shift & Attendance", f"{PORTAL}/app/shift-and-attendance"],
        ["Performance", f"{PORTAL}/app/performance"],
        ["Recruitment", f"{PORTAL}/app/recruitment"],
        ["Employee self-service", "Frappe HR mobile app (point to erp.vsjailabs.in)"],
    ], widths=[5, 11], code_cols=(1,))
    heading(doc, "3.2 Login", 2)
    bullet(doc, "Open the portal at " + PORTAL + " in a modern browser (Chrome/Firefox/Edge/Safari).", num=True)
    bullet(doc, "Enter your username (email) and password. New users receive a welcome email with a link to set their own password.", num=True)
    bullet(doc, "Each staff member needs an Employee record linked to their User account to use self-service features.", num=True)
    doc.add_page_break()

    # ── 4. ROLES ──
    heading(doc, "4. Roles & Permissions", 1)
    body(doc, "HRMS access is controlled by roles assigned to each User (Settings > User > Roles). Assign the minimum role needed.")
    make_table(doc, ["Role", "Scope", "Typical User"], [
        ["HR Manager", "Full HR setup + approval authority (leave, payroll, claims)", "HR Head / Director"],
        ["HR User", "Day-to-day HR operations (data entry, attendance, applications)", "HR Executive"],
        ["Employee", "Self-service: own leave, attendance, claims, payslips", "All staff"],
        ["Leave Approver", "Approve leave for reportees (assigned per employee)", "Reporting Managers"],
        ["Expense Approver", "Approve expense claims for reportees", "Reporting Managers / Finance"],
    ], widths=[3.5, 9, 3.5])
    note(doc, "Tip", "Set each employee's Reporting Manager, Leave Approver, and Expense Approver on their Employee record so approval routing and notification emails work automatically.")
    doc.add_page_break()

    # ── 5. SETUP ──
    heading(doc, "5. One-Time Setup (Prerequisites)", 1)
    note(doc, "Important", "HRMS is dependency-chained: Employee → Attendance → Leave → Payroll each build on the prior layer's masters. Complete this setup in order — skipping a master (e.g. Holiday List) silently breaks downstream leave balances and payroll day calculations.", color=ORANGE)
    heading(doc, "5.1 Setup Order", 2)
    make_table(doc, ["#", "Step", "Path", "Notes"], [
        ["1", "Company", "/app/company", "Exists from ERPNext; set default payroll payable account"],
        ["2", "HR Settings", "/app/hr-settings", "Working hours, retirement age, employee naming, reminders"],
        ["3", "Department", "/app/department", "Org structure"],
        ["4", "Designation", "/app/designation", "Job titles"],
        ["5", "Branch", "/app/branch", "Locations (optional)"],
        ["6", "Employment Type", "/app/employment-type", "Full-time, Contract, Intern, etc."],
        ["7", "Employee Grade", "/app/employee-grade", "Pay grades (optional)"],
        ["8", "Holiday List", "/app/holiday-list", "REQUIRED — weekly offs + public holidays, per year"],
    ], widths=[1, 4, 5, 6], code_cols=(2,))
    doc.add_page_break()

    # ── 6. EMPLOYEE ──
    heading(doc, "6. Employee Management", 1)
    heading(doc, "6.1 Create an Employee", 2)
    bullet(doc, "Go to HR > Employee > New (/app/employee/new).", num=True)
    bullet(doc, "Fill: Name, Gender, Date of Birth, Date of Joining, Company, Department, Designation, Employment Type.", num=True)
    bullet(doc, "Link the employee to a User account (Preferred Email + 'Create User') so they get self-service access.", num=True)
    bullet(doc, "Set Reporting Manager, Leave Approver, Expense Approver, and Holiday List for routing and calculations.", num=True)
    bullet(doc, "Save. The employee can now appear in attendance, leave, and payroll.", num=True)
    heading(doc, "6.2 Bulk Import", 2)
    body(doc, "To load many employees at once: Employee list > Menu (⋯) > Import > download the template, fill it, and upload. Map columns and run the import.")
    heading(doc, "6.3 Onboarding", 2)
    body(doc, "HR > Employee Onboarding lets you define a reusable checklist (IT setup, document collection, induction) that auto-creates tasks when a new hire joins, using an Employee Onboarding Template.")
    doc.add_page_break()

    # ── 7. ATTENDANCE ──
    heading(doc, "7. Attendance & Shift Management", 1)
    make_table(doc, ["Task", "Path", "Notes"], [
        ["Mark attendance (single)", "/app/attendance/new", "One day per employee"],
        ["Bulk mark attendance", "HR workspace > Mark Attendance", "Select dates + employees"],
        ["Check-in / Check-out", "/app/employee-checkin", "Mobile app, biometric, or manual"],
        ["Shift Type", "/app/shift-type", "Defines timings + auto-attendance from check-ins"],
        ["Shift Assignment", "/app/shift-assignment", "Assign shifts to employees"],
        ["Attendance Request", "/app/attendance-request", "On-duty / WFH / regularization"],
    ], widths=[4.5, 5.5, 6], code_cols=(1,))
    note(doc, "Auto-attendance", "Configure a Shift Type with 'Enable Auto Attendance' to generate Attendance automatically from Employee Checkin records (biometric or mobile).")
    doc.add_page_break()

    # ── 8. LEAVE ──
    heading(doc, "8. Leave Management", 1)
    heading(doc, "8.1 Setup", 2)
    make_table(doc, ["#", "Step", "Path", "Purpose"], [
        ["1", "Leave Type", "/app/leave-type", "Casual, Sick, Earned — max days, carry-forward, LWP"],
        ["2", "Leave Period", "/app/leave-period", "The cycle, e.g. FY 2026-27"],
        ["3", "Leave Policy", "/app/leave-policy", "Annual allocation per leave type"],
        ["4", "Leave Policy Assignment", "/app/leave-policy-assignment", "Grant the policy to employees/grades"],
        ["5", "Leave Allocation", "/app/leave-allocation", "Opening balances (auto from policy or manual)"],
    ], widths=[1, 5, 5.5, 5], code_cols=(2,))
    heading(doc, "8.2 Applying for Leave", 2)
    bullet(doc, "Employee: HR > Leave Application > New (or via mobile app). Select leave type and dates.", num=True)
    bullet(doc, "The Leave Approver / reporting manager is notified by email and approves or rejects.", num=True)
    bullet(doc, "On approval, the leave balance is deducted automatically.", num=True)
    doc.add_page_break()

    # ── 9. PAYROLL ──
    heading(doc, "9. Payroll", 1)
    make_table(doc, ["#", "Step", "Path", "Purpose"], [
        ["1", "Payroll Period", "/app/payroll-period", "Defines the tax/period window"],
        ["2", "Salary Component", "/app/salary-component", "Earnings (Basic, HRA) & Deductions (PF, TDS)"],
        ["3", "Salary Structure", "/app/salary-structure", "Assemble components with formulas"],
        ["4", "Salary Structure Assignment", "/app/salary-structure-assignment", "Attach structure + base to each employee"],
        ["5", "Payroll Entry", "/app/payroll-entry", "Run for a month/department → bulk Salary Slips"],
        ["6", "Salary Slip", "/app/salary-slip", "Individual payslip; employees view their own"],
    ], widths=[1, 5, 5.5, 5], code_cols=(2,))
    heading(doc, "9.1 Running Monthly Payroll", 2)
    bullet(doc, "Ensure every employee has a submitted Salary Structure Assignment effective for the month.", num=True)
    bullet(doc, "Create a Payroll Entry: select company, payroll frequency, dates, and department/branch filters.", num=True)
    bullet(doc, "Click 'Create Salary Slips', then 'Submit Salary Slips'. Review and submit.", num=True)
    bullet(doc, "Generate the Bank Entry / Journal Entry to record payment; payslips can be emailed to employees.", num=True)
    note(doc, "India", "For TDS/Income Tax, configure deduction components with formulas and collect Employee Tax Exemption Declarations and Proof Submissions.")
    doc.add_page_break()

    # ── 10-13 combined ──
    heading(doc, "10. Expense Claims", 1)
    body(doc, "Employees submit out-of-pocket expenses for reimbursement.")
    bullet(doc, "Employee: HR > Expense Claim > New — add expense lines, attach receipts, submit.")
    bullet(doc, "Expense Approver reviews and approves; the approved amount is reimbursed via Payment Entry or through payroll.")
    bullet(doc, "Configure Expense Claim Types (/app/expense-claim-type) for categorisation and default accounts.")

    heading(doc, "11. Performance Management", 1)
    bullet(doc, "Define an Appraisal Cycle (/app/appraisal-cycle) for the review period.")
    bullet(doc, "Use Appraisal Templates to set KRAs/goals with weightage.")
    bullet(doc, "Employees self-rate; managers review and finalise the Appraisal score and feedback.")

    heading(doc, "12. Recruitment", 1)
    bullet(doc, "Create a Job Opening (/app/job-opening) tied to a Designation and Department.")
    bullet(doc, "Track candidates as Job Applicants; record interview feedback.")
    bullet(doc, "Issue a Job Offer; on acceptance, convert the applicant into an Employee.")

    heading(doc, "13. Employee Lifecycle", 1)
    bullet(doc, "Employee Promotion (/app/employee-promotion) — change designation/grade with effective date.")
    bullet(doc, "Employee Transfer (/app/employee-transfer) — move between departments/branches/companies.")
    bullet(doc, "Employee Separation (/app/employee-separation) — structured exit checklist; set relieving date.")
    doc.add_page_break()

    # ── 14. REPORTS ──
    heading(doc, "14. Reports & Dashboards", 1)
    body(doc, "Built-in HR reports are available from the HR and Payroll workspaces:")
    make_table(doc, ["Report", "Shows"], [
        ["Employee Leave Balance", "Remaining leave per employee by type"],
        ["Monthly Attendance Sheet", "Day-by-day attendance grid for a month"],
        ["Salary Register", "Consolidated payroll for a period"],
        ["Employee Birthday", "Upcoming birthdays"],
        ["Employee Information", "Headcount, demographics, tenure"],
    ], widths=[5.5, 10.5])
    doc.add_page_break()

    # ── 15. MOBILE ──
    heading(doc, "15. Mobile App (Frappe HR)", 1)
    bullet(doc, "Install 'Frappe HR' from the Apple App Store or Google Play.")
    bullet(doc, "On first launch, enter the server URL: erp.vsjailabs.in and sign in with your ERPNext credentials.")
    bullet(doc, "Employees can: check in/out (with geolocation), apply for leave, submit expense claims, view payslips, and approve requests (managers).")
    doc.add_page_break()

    # ── 16. FIRST RUN ──
    heading(doc, "16. Recommended First-Run Path", 1)
    body(doc, "To get value quickly, follow this sequence end-to-end on one employee before rolling out:")
    for i, t in enumerate([
        "Complete HR Settings and create a Holiday List for the current year.",
        "Create a few Departments and Designations.",
        "Add one Employee and link them to a User (welcome email goes out automatically).",
        "Set up one Leave Type + Leave Policy and allocate leave.",
        "Have the employee apply for leave; approve it (tests routing + email).",
        "Build one Salary Structure, assign it, and run a Payroll Entry for that employee.",
        "Review the generated Salary Slip and email it.",
    ], 1):
        bullet(doc, t, num=True)
    doc.add_page_break()

    # ── 17. EMAIL ──
    heading(doc, "17. Email & Notifications", 1)
    body(doc, "Outgoing email is configured on this instance, so HR notifications are delivered automatically.")
    make_table(doc, ["Setting", "Value"], [
        ["SMTP relay", "Zoho (smtp.zoho.in, port 465, SSL)"],
        ["Sender (From)", "admin@vsjailabs.com"],
        ["Welcome email on new User", "Enabled by default (link to set password)"],
        ["Leave/expense approval emails", "Sent to the assigned approver"],
        ["Delivery", "Queue-based (Email Queue), flushed by the scheduler (enabled)"],
    ], widths=[6, 10], code_cols=(1,))
    note(doc, "Note", "Outgoing mail is sent from admin@vsjailabs.com (not noreply@), and the scheduler must remain enabled for queued mail to flush. First emails to a recipient may land in Spam until the sender reputation warms up.")
    doc.add_page_break()

    # ── 18. TROUBLESHOOTING ──
    heading(doc, "18. Troubleshooting & FAQ", 1)
    make_table(doc, ["Symptom", "Likely Cause", "Resolution"], [
        ["Leave balance is zero / can't apply", "No Leave Allocation or Holiday List", "Create Leave Allocation (or Policy Assignment) and assign a Holiday List"],
        ["Payroll skips an employee", "No active Salary Structure Assignment", "Create & submit an assignment effective for the period"],
        ["New user didn't get welcome email", "'Send Welcome Email' unchecked, or in Spam", "Re-check the box / resend; check Spam folder"],
        ["Attendance not auto-created", "Shift auto-attendance off or no check-ins", "Enable Auto Attendance on Shift Type; verify Employee Checkin records"],
        ["Account locked after wrong password", "Brute-force protection", "Wait a few minutes or have HR/Admin reset"],
        ["Employee can't see self-service", "User not linked / no Employee role", "Link User on Employee record; ensure Employee role"],
    ], widths=[5, 5, 6])
    doc.add_page_break()

    # ── 19. SUPPORT ──
    heading(doc, "19. Support & Escalation", 1)
    make_table(doc, ["Level", "Contact", "Scope"], [
        ["Level 1 — HR", "HR team (internal)", "Day-to-day HR queries, data entry, approvals"],
        ["Level 2 — System Admin", AUTHOR + " / IT", "Configuration, roles, payroll setup, email"],
        ["Reference", "frappe.io/hr docs", "Official Frappe HR documentation"],
    ], widths=[4.5, 5.5, 6])

    # Footer block
    doc.add_paragraph()
    fb = doc.add_table(rows=1, cols=1); fb.alignment = WD_TABLE_ALIGNMENT.CENTER
    shade(fb.rows[0].cells[0], NAVY_HEX)
    c = fb.rows[0].cells[0]; c.text = ""
    p = c.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f"Prepared by {AUTHOR}  |  {COMPANY}\n{DOC_TITLE} v{VERSION}  |  {TODAY}  |  {CLASSIFICATION}")
    r.font.name = "Arial"; r.font.size = Pt(9); r.font.bold = True; r.font.color.rgb = WHITE

    out = "/Users/tpe/VSJWORK/erpnext-vsjailabs/docs/VSJ_HRMS_Functional_Guide_v1.0.docx"
    doc.save(out)
    print("SAVED", out)

if __name__ == "__main__":
    generate()
